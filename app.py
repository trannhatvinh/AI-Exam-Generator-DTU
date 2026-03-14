import os
import json
import random
import string
import re
import numpy as np
import time
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, jsonify, flash

from docx import Document
import pdfplumber
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from sentence_transformers import SentenceTransformer

from nltk.translate.bleu_score import sentence_bleu
from rouge_score import rouge_scorer

from groq import Groq

from ai_pipeline import *

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False

UPLOAD_FOLDER = "topics"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
MODEL_NAME = "llama-3.1-8b-instant"
GROQ_API_KEY = "YOUR_API_KEY"

client = Groq(api_key=GROQ_API_KEY)
embedder = SentenceTransformer("paraphrase-multilingual-MiniLM-L12-v2")

app = Flask(__name__)
app.secret_key = 'test-secret'

def gen_batch_questions(context, qtype, bloom_level, difficulty, used_questions=None, batch_size=20):
    used_questions = used_questions or set()
    diff_text = difficulty_prompt_text(difficulty)
    if qtype == 'essay':
        prompt = f"""{diff_text}
Sinh {batch_size} câu hỏi essay (tự luận ngắn) đa dạng chủ đề, không trùng nhau, MỖI CÂU PHẢI LUÔN có answer trả lời ngắn gọn, đầy đủ, chỉ trả về mảng JSON đúng [{batch_size} phần tử], không chú thích, không bỏ trống trường answer.
[
  {{"type": "essay", "question": "...", "answer": "..."}},
  ...
]
Context:
{context[:1800]}
"""
    elif qtype == 'mcq':
        prompt = f"""{diff_text}

Sinh {batch_size} câu hỏi trắc nghiệm (MCQ).

YÊU CẦU BẮT BUỘC:

- Mỗi câu có 4 lựa chọn.
- options phải là NỘI DUNG đáp án, KHÔNG dùng A, B, C, D.
- KHÔNG dùng nhãn "A.", "B.", "C.", "D."
- Các lựa chọn phải cùng loại và hợp lý.
- Không dùng placeholder như "A","B","C","D".
- answer là chỉ số đúng (0-3).

Chỉ trả về JSON array đúng định dạng:

[
  {{
    "type": "mcq",
    "question": "...",
    "options": [
      "đáp án 1",
      "đáp án 2",
      "đáp án 3",
      "đáp án 4"
    ],
    "answer": 0
  }}
]

Context:
{context[:1800]}
"""
    log(f"Đang gọi AI (batch: {qtype}, size: {batch_size})...")
    start = time.time()
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.35,
        max_tokens=1700
    )
    elapsed = time.time() - start
    log(f"...xong AI (batch {qtype}, {elapsed:.1f} giây).")
    text = response.choices[0].message.content
    text = text.replace("```json", "").replace("```", "")
    try:
        batch = json.loads(text)
        result = []
        for q in batch:
            if (
                isinstance(q, dict)
                and q.get("type") == qtype
                and q.get("question")
            ):
                # Đảm bảo essay luôn có answer đầy đủ
                if qtype == "essay":
                    ans = q.get("answer", "")
                    if not ans or len(str(ans).strip()) < 6:
                        q["answer"] = "[Chưa có đáp án]"
                result.append(q)
        log(f"Sinh về được {len(result)} câu {qtype}.")
        return result
    except Exception:
        matches = re.findall(r"\{[\s\S]+?\}", text)
        result = []
        for m in matches:
            try:
                q = json.loads(m)
                if (
                    q.get("type") == qtype
                    and q.get("question")
                ):
                    if qtype == "essay":
                        ans = q.get("answer", "")
                        if not ans or len(str(ans).strip()) < 6:
                            q["answer"] = "[Chưa có đáp án]"
                    result.append(q)
            except:
                continue
        log(f"Sinh được (fallback parsers) {len(result)} câu {qtype}.")
        return result

def filter_exact_uniques(question_list, quota=None):
    seen = set()
    filtered = []
    for q in question_list:
        text = q.get("question", "").strip()
        if text and text not in seen:
            filtered.append(q)
            seen.add(text)
        if quota and len(filtered) >= quota:
            break
    return filtered

def generate_questions_from_clusters(clusters, num_mcq, num_short, difficulty, existing_questions=None):
    batch_size = max(num_mcq, num_short, 10)
    existed_texts = set(q.get("question", "").strip() for q in existing_questions) if existing_questions else set()
    questions = []
    quota_mcq = num_mcq
    quota_essay = num_short
    log(f"Tổng cluster: {len(clusters)}, sinh {num_mcq} TN, {num_short} TL, batch_size: {batch_size}...")
    for inx, c in enumerate(clusters):
        if quota_mcq > 0:
            batch = gen_batch_questions(c, 'mcq', "AAA", difficulty, used_questions=existed_texts, batch_size=quota_mcq)
            for q in batch:
                text = q.get("question", "").strip()
                if text and text not in existed_texts:
                    q["difficulty"] = difficulty
                    questions.append(q)
                    existed_texts.add(text)
            quota_mcq = num_mcq - len([q for q in questions if q["type"] == "mcq"])
            log(f"MCQ còn thiếu: {quota_mcq}")
            if quota_mcq <= 0:
                break
    for inx, c in enumerate(clusters):
        if quota_essay > 0:
            batch = gen_batch_questions(c, 'essay', "AAA", difficulty, used_questions=existed_texts, batch_size=quota_essay)
            for q in batch:
                text = q.get("question", "").strip()
                if text and text not in existed_texts:
                    q["difficulty"] = difficulty
                    questions.append(q)
                    existed_texts.add(text)
            quota_essay = num_short - len([q for q in questions if q["type"] == "essay"])
            log(f"Essay còn thiếu: {quota_essay}")
            if quota_essay <= 0:
                break
    mcq_unique = filter_exact_uniques([q for q in questions if q["type"] == "mcq"], quota=num_mcq)
    essay_unique = filter_exact_uniques([q for q in questions if q["type"] == "essay"], quota=num_short)
    log(f"==> Tổng sinh xong {len(mcq_unique)} TN và {len(essay_unique)} TL.")
    return {"questions": mcq_unique + essay_unique}

def export_official_exam_format(
        topic_path, selected_mcq, selected_essay, num_mcq, num_essay,
        points_mcq, points_essay, file_name="exam_official.docx",
        subject=".......", code="A", hoc_ky=".......", nam_hoc=".......",
        lop=".......", time_str="75 phút", teacher="......................", leader="......................"
):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(1.5), Cm(1.5)
    section.left_margin, section.right_margin = Cm(2.0), Cm(1.5)

    # ===== FIX HEADER: Ép cứng độ rộng từng ô để cột 3 thật sự nhỏ =====
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.allow_autofit = False
    # Nới chiều cao hàng (Row Height) - Đây là điểm mấu chốt để bảng không bị sát
    table.rows[0].height = Cm(2.5)  # Thiết lập chiều cao tối thiểu là 2.5cm
    # Thiết lập độ rộng cột (5.0 - 11.5 - 1.5 = 18.0cm tổng cộng)
    widths = [Cm(5.0), Cm(11.5), Cm(1.5)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
        table.cell(0, i).width = width  # Ép độ rộng cho từng ô

    # Cột 1: Đơn vị [cite: 1]
    c1 = table.cell(0, 0).paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.add_run("ĐẠI HỌC DUY TÂN\n").bold = True
    c1.add_run("Trường KHMT\n").bold = True
    c1.add_run("Khoa CNTT").bold = True

    # Cột 2: Tiêu đề thi [cite: 1]
    c2 = table.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = c2.add_run("ĐỀ THI KẾT THÚC HỌC PHẦN\n")
    run_t.bold = True
    c2.add_run(f"Môn: {subject}\n")
    c2.add_run(f"Khối lớp: {lop}  Học kỳ: {hoc_ky}  Năm học: {nam_hoc}\n")
    c2.add_run(f"Thời gian làm bài: {time_str}")


    # Cột 3: Mã đề (NHỎ LẠI) [cite: 1]
    c3 = table.cell(0, 2).paragraphs[0]
    c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c3.add_run("ĐỀ\n").bold = True
    c3.add_run(code).bold = True

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- Thông tin SV ---
    p_sv = doc.add_paragraph()
    # Tăng khoảng cách phía trên (với Header) và phía dưới (với Box Lưu ý)
    p_sv.paragraph_format.space_before = Pt(18)
    p_sv.paragraph_format.space_after = Pt(12)
    p_sv.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_sv.paragraph_format.line_spacing = Pt(18)
    p_sv.add_run(f"Họ tên sinh viên: {'.' * 55} MSSV: {'.' * 25}")

    # --- Box Lưu ý ---
    note_table = doc.add_table(rows=1, cols=1)
    note_table.style = 'Table Grid'
    # Đặt độ rộng bảng lưu ý chiếm gần hết chiều ngang để không bị méo
    note_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_table.width = Cm(17.0)

    mcq_p_each = round(points_mcq / num_mcq, 3) if num_mcq > 0 else 0
    essay_p_each = round(points_essay / num_essay, 2) if num_essay > 0 else 0

    note_cell = note_table.cell(0, 0)
    # Thiết lập padding cho ô (khoảng cách chữ với viền bảng)
    note_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    note_p = note_cell.paragraphs[0]
    # Căn chỉnh dòng trong Box Lưu ý: 0pt before/after, Exactly 18pt để các dòng lưu ý không dính nhau
    note_p.paragraph_format.space_before = Pt(6)
    note_p.paragraph_format.space_after = Pt(6)
    note_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    note_p.paragraph_format.line_spacing = Pt(18)

    note_p.add_run("LƯU Ý:\n").bold = True
    note_p.add_run(f"• Đề thi gồm 2 phần:\n")
    note_p.add_run(f"   + Phần 1: Trắc nghiệm ({num_mcq} câu × {mcq_p_each:g} điểm = {points_mcq:g} điểm)\n")
    note_p.add_run(f"   + Phần 2: Câu tự luận ({num_essay} câu × {essay_p_each:g} điểm = {points_essay:g} điểm)\n")
    note_p.add_run("- Sinh viên không làm bài trực tiếp vào đề thi.\n")
    note_p.add_run(f"- {num_mcq} câu trắc nghiệm: sinh viên làm trên phiếu trả lời.\n")
    note_p.add_run(f"- {num_essay} câu tự luận: sinh viên viết trực tiếp vào phần trả lời dưới đây.")

    # Phần 1: Trắc nghiệm
    p1_header = doc.add_paragraph()
    p1_header.paragraph_format.space_before = Pt(12)
    p1_header.add_run(f"Phần 1. Trắc nghiệm ({points_mcq:g} điểm)").bold = True
    p1_header.paragraph_format.space_after = Pt(0)
    for i, q in enumerate(selected_mcq, 1):
        q_p = doc.add_paragraph()
        # THIẾT LẬP SPACING THEO YÊU CẦU
        q_p.paragraph_format.space_before = Pt(0)
        q_p.paragraph_format.space_after = Pt(0)
        q_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        q_p.paragraph_format.line_spacing = Pt(18)

        q_p.add_run(f"Câu {i} ({mcq_p_each:g} điểm): ").bold = True
        q_p.add_run(q["question"])

        for j, opt in enumerate(q["options"]):
            p_opt = doc.add_paragraph()
            p_opt.paragraph_format.space_before = Pt(0)
            p_opt.paragraph_format.space_after = Pt(0)
            p_opt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p_opt.paragraph_format.line_spacing = Pt(18)
            p_opt.paragraph_format.left_indent = Cm(0.75)
            p_opt.add_run(f"{chr(65 + j)}. {opt}")

    # Phần 2: Tự luận (Đã sửa Spacing & Line Spacing 18pt)
    p2_header = doc.add_paragraph()
    p2_header.paragraph_format.space_before = Pt(12)
    p2_header.add_run(f"Phần 2. Tự luận ({points_essay:g} điểm)").bold = True
    p2_header.paragraph_format.space_after = Pt(0)
    for i, q in enumerate(selected_essay, num_mcq + 1):
        q_p = doc.add_paragraph()
        q_p.paragraph_format.space_before = Pt(6)  # Giữ một chút khoảng cách giữa các câu
        q_p.paragraph_format.space_after = Pt(0)
        q_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        q_p.paragraph_format.line_spacing = Pt(18)

        q_p.add_run(f"Câu {i} ({essay_p_each:g} điểm): ").bold = True
        q_p.add_run(q["question"])

        # Thêm các dòng kẻ chấm
        line = doc.add_paragraph("." * 105)
        line.paragraph_format.space_before = Pt(0)
        line.paragraph_format.space_after = Pt(0)
        line.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        line.paragraph_format.line_spacing = Pt(18)

    # --- Dòng ghi chú in nghiêng nằm giữa ---
    p_note_end = doc.add_paragraph()
    p_note_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note_end.paragraph_format.space_before = Pt(12)  # Khoảng cách với phần nội dung trên
    p_note_end.paragraph_format.space_after = Pt(12)
    run_note = p_note_end.add_run("(Đề thi không sử dụng tài liệu cán bộ coi thi không giải thích gì thêm)")
    run_note.italic = True

    # Chữ ký [cite: 13, 14]
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.cell(0, 0).paragraphs[0].add_run("Tổ trưởng bộ môn").bold = True
    sig_table.cell(0, 1).paragraphs[0].add_run("Giảng viên ra đề").bold = True
    sig_table.cell(1, 0).paragraphs[0].add_run(f"\n\n\n{leader}")
    sig_table.cell(1, 1).paragraphs[0].add_run(f"\n\n\n{teacher}")
    for row in sig_table.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(os.path.join(topic_path, "exam", file_name))
    return file_name


def export_official_answer(
        topic_path, selected_mcq, selected_essay, num_mcq, num_essay,
        points_mcq, points_essay, file_name="answer_official.docx",
        subject=".......", code="A", hoc_ky=".......", nam_hoc=".......",
        lop=".......", time_str="...", teacher="......................", leader="......................"
):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(1.5), Cm(1.5)
    section.left_margin, section.right_margin = Cm(2.0), Cm(1.5)

    # ===== FIX HEADER: Ép cứng độ rộng từng ô để cột 3 thật sự nhỏ =====
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.allow_autofit = False
    # Nới chiều cao hàng (Row Height) - Đây là điểm mấu chốt để bảng không bị sát
    table.rows[0].height = Cm(2.5)  # Thiết lập chiều cao tối thiểu là 2.5cm
    # Thiết lập độ rộng cột (5.0 - 11.5 - 1.5 = 18.0cm tổng cộng)
    widths = [Cm(5.0), Cm(11.5), Cm(1.5)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
        table.cell(0, i).width = width  # Ép độ rộng cho từng ô

    # Cột 1: Đơn vị [cite: 1]
    c1 = table.cell(0, 0).paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.add_run("ĐẠI HỌC DUY TÂN\n").bold = True
    c1.add_run("Trường KHMT\n").bold = True
    c1.add_run("Khoa CNTT").bold = True

    # Cột 2: Tiêu đề thi [cite: 1]
    c2 = table.cell(0, 1).paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = c2.add_run("ĐÁP ÁN ĐỀ THI KẾT THÚC HỌC PHẦN\n")
    run_t.bold = True
    c2.add_run(f"Môn: {subject}\n")
    c2.add_run(f"Khối lớp: {lop}  Học kỳ: {hoc_ky}  Năm học: {nam_hoc}\n")

    # Cột 3: Mã đề (NHỎ LẠI) [cite: 1]
    c3 = table.cell(0, 2).paragraphs[0]
    c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c3.add_run("ĐỀ\n").bold = True
    c3.add_run(code).bold = True

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Đáp án Trắc nghiệm (Bảng ngang) [cite: 18]
    doc.add_paragraph().add_run(f"\nPhần 1. Đáp án phần trắc nghiệm ({points_mcq:g} điểm)").bold = True
    ans_table = doc.add_table(rows=2, cols=len(selected_mcq))
    ans_table.style = 'Table Grid'
    for i, q in enumerate(selected_mcq):
        ans_table.cell(0, i).text = str(i + 1)
        ans_table.cell(1, i).text = chr(65 + q["answer"])
        for row in range(2):
            ans_table.cell(row, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Đáp án Tự luận [cite: 20]
    doc.add_paragraph().add_run(f"\nPhần 2. Đáp án phần Tự luận ({points_essay:g} điểm)").bold = True
    for i, q in enumerate(selected_essay, num_mcq + 1):
        p = doc.add_paragraph()
        p.add_run(f"Câu {i}: ").bold = True
        p.add_run(q.get("answer", "............"))

    # Chữ ký [cite: 22, 23]
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.cell(0, 0).paragraphs[0].add_run("Tổ trưởng bộ môn").bold = True
    sig_table.cell(0, 1).paragraphs[0].add_run("Giảng viên ra đề").bold = True
    sig_table.cell(1, 0).paragraphs[0].add_run(f"\n\n\n{leader}")
    sig_table.cell(1, 1).paragraphs[0].add_run(f"\n\n\n{teacher}")
    for row in sig_table.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(os.path.join(topic_path, "exam", file_name))
    return file_name

def generate_word_exams_from_selected(topic_path, selected_mcq, selected_essay, points_mcq, points_essay, code):
    exam_path = os.path.join(topic_path, "exam")
    os.makedirs(exam_path, exist_ok=True)
    doc = Document()
    doc.add_heading(f"ĐỀ THI - MÃ {code}", level=1)
    idx = 1

    total_mcq = len(selected_mcq)
    total_essay = len(selected_essay)
    point_each_mcq = round(points_mcq / max(total_mcq, 1), 2)
    point_each_essay = round(points_essay / max(total_essay, 1), 2)

    import re
    # Add MCQ
    for q in selected_mcq:
        p1 = doc.add_paragraph()
        p1.add_run(f"{idx}. ({point_each_mcq} Point) ").bold = True
        p2 = doc.add_paragraph()
        p2.add_run(q["question"]).bold = True
        options = list(zip(q["options"], range(len(q["options"]))))
        # Đảm bảo không xáo trộn options so với file official
        for j, (opt, orig) in enumerate(options):
            s = opt.strip()
            if re.match(r"^\*?[A-D]\.", s):
                text = s.lstrip("*")
            else:
                letter = chr(65 + j)
                text = f"{letter}. {opt}"
            if orig == q["answer"]:
                text = "*" + text
            doc.add_paragraph(text)
        idx += 1

    # Add ESSAY
    for q in selected_essay:
        p1 = doc.add_paragraph()
        p1.add_run(f"{idx}. ({point_each_essay} Point) ").bold = True
        p2 = doc.add_paragraph()
        p2.add_run(q["question"]).bold = True
        doc.add_paragraph("Trả lời: ................................................")
        idx += 1

    filename = f"exam_{code}.docx"
    path = os.path.join(exam_path, filename)
    doc.save(path)
    return filename

@app.route('/manage_topic_files', methods=['GET', 'POST'])
def manage_topic_files():
    all_topics = sorted(os.listdir(UPLOAD_FOLDER))
    current_topic = request.args.get('selected_topic') or None
    files = []
    if request.method == 'POST':
        t = request.form.get('topic') or request.form.get('new_topic', '').strip()
        if request.form['action'] == 'add_topic' and t:
            topic_folder = os.path.join(UPLOAD_FOLDER, t)
            os.makedirs(os.path.join(topic_folder, "uploads"), exist_ok=True)
            flash(f'Đã tạo chủ đề mới: {t}', 'success')
            return redirect(url_for('manage_topic_files', selected_topic=t))
        elif request.form['action'] == 'upload_file' and t:
            file = request.files.get('file')
            if file:
                upload_dir = os.path.join(UPLOAD_FOLDER, t, "uploads")
                os.makedirs(upload_dir, exist_ok=True)
                file.save(os.path.join(upload_dir, file.filename))
                current_topic = t
        elif request.form['action'] == 'delete_file' and t:
            fn = request.form.get('file_to_delete')
            if fn:
                upload_dir = os.path.join(UPLOAD_FOLDER, t, "uploads")
                fp = os.path.join(upload_dir, fn)
                if os.path.exists(fp):
                    os.remove(fp)
                current_topic = t
    if current_topic:
        upload_dir = os.path.join(UPLOAD_FOLDER, current_topic, "uploads")
        if os.path.exists(upload_dir):
            files = os.listdir(upload_dir)
        else:
            files = []
    return render_template('form_upload_bank.html',
                          all_topics=all_topics,
                          current_topic=current_topic,
                          files=files)

@app.route('/api/get_topic_files/<topic>')
def api_get_topic_files(topic):
    upload_dir = os.path.join(UPLOAD_FOLDER, topic, "uploads")
    if not os.path.isdir(upload_dir):
        files = []
    else:
        files = [f for f in os.listdir(upload_dir)]
    return render_template('partials/_file_list.html', files=files, current_topic=topic)

@app.route('/api/delete_topic_file/<topic>/<filename>', methods=['POST'])
def api_delete_topic_file(topic, filename):
    upload_dir = os.path.join(UPLOAD_FOLDER, topic, "uploads")
    fpath = os.path.join(upload_dir, filename)
    if os.path.exists(fpath):
        os.remove(fpath)
        return jsonify(success=True)
    else:
        return jsonify(success=False)

@app.route("/generate_bank", methods=["GET", "POST"])
def generate_bank():
    success = False
    q_count = 0
    metrics = ""
    topics = [t for t in os.listdir(UPLOAD_FOLDER) if os.path.isdir(os.path.join(UPLOAD_FOLDER, t))]
    current_topic = None
    if request.method == "POST":
        log("Bắt đầu generate_bank...")
        t0 = time.time()
        topic = request.form["topic"]
        num_mcq = int(request.form["num_mcq"])
        num_short = int(request.form["num_short"])
        difficulty = request.form["difficulty"]
        current_topic = topic
        topic_path = os.path.join(UPLOAD_FOLDER, topic)
        all_texts = load_all_texts_from_topic(topic_path)
        clusters = select_clusters_from_texts(all_texts, cluster_size=1800, max_clusters=8)
        bank_file = os.path.join(topic_path, "question_bank.json")
        existing_questions = []
        if os.path.exists(bank_file):
            with open(bank_file, encoding="utf-8") as f:
                data = json.load(f)
            existing_questions = data.get("questions", [])
        data = generate_questions_from_clusters(clusters, num_mcq, num_short, difficulty, existing_questions=existing_questions)
        questions = normalize_questions(data.get("questions", []))
        all_questions = existing_questions + questions
        unique_questions = []
        texts = set()
        for q in all_questions:
            text_q = q.get("question", "").strip()
            if text_q not in texts and text_q:
                texts.add(text_q)
                unique_questions.append(q)
        data = {"questions": unique_questions}
        with open(bank_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        context = "\n".join(all_texts)
        bleu, rouge, acc = compute_metrics(questions, context)
        metrics = f"""Topic: {topic}
Questions: {len(questions)}
BLEU: {bleu:.4f}
ROUGE: {rouge:.4f}
Accuracy: {acc:.4f}
"""
        with open(os.path.join(topic_path, "metrics.txt"), "w", encoding="utf-8") as f:
            f.write(metrics)
        q_count = len(questions)
        success = True
        log(f"Hoàn tất generate_bank toàn bộ ({int(time.time()-t0)} giây).")
    return render_template("form_bank.html", success=success, q_count=q_count, metrics=metrics, topics=topics, current_topic=current_topic)

@app.route("/generate_exam", methods=["GET", "POST"])
def generate_exam():
    topics = os.listdir(UPLOAD_FOLDER)
    files = []
    current_topic = None
    warnings = []
    if request.method == "POST":
        topic = request.form["topic"]
        current_topic = topic
        num_mcq = int(request.form.get("num_mcq", 10))
        num_essay = int(request.form.get("num_essay", 2))
        points_mcq = float(request.form.get("points_mcq", 2.5))
        points_essay = float(request.form.get("points_essay", 7.5))
        versions = int(request.form.get("versions", 2))
        mode = request.form.get("distribute_mode", "auto")
        percent_easy = int(request.form.get("percent_easy", 40))
        percent_medium = int(request.form.get("percent_medium", 40))
        percent_hard = int(request.form.get("percent_hard", 20))
        topic_path = os.path.join(UPLOAD_FOLDER, topic)
        with open(os.path.join(topic_path, "question_bank.json"), encoding="utf-8") as f:
            data = json.load(f)["questions"]
        mcq = [q for q in data if q["type"] == "mcq"]
        essay = [q for q in data if q["type"] == "essay"]
        def calc_distribute(n, p_easy, p_medium, p_hard):
            c_easy = round(n * p_easy / 100)
            c_medium = round(n * p_medium / 100)
            c_hard = n - c_easy - c_medium
            arr = [("Easy", c_easy), ("Medium", c_medium), ("Hard", c_hard)]
            while sum(x[1] for x in arr) != n:
                diff = n - sum(x[1] for x in arr)
                for i in range(len(arr)):
                    if diff == 0:
                        break
                    arr[i] = (arr[i][0], arr[i][1] + (1 if diff > 0 else -1))
                    diff = n - sum(x[1] for x in arr)
            return dict(arr)
        if mode == "manual":
            counts_mcq = calc_distribute(num_mcq, percent_easy, percent_medium, percent_hard)
            counts_essay = calc_distribute(num_essay, percent_easy, percent_medium, percent_hard)
        else:
            def random_counts(pool, n):
                lvls = ["Easy", "Medium", "Hard"]
                l = [q.get("difficulty", "Medium") for q in pool]
                idx_pool = list(range(len(pool)))
                pick = random.sample(idx_pool, min(n, len(pool)))
                picked_lvls = [pool[i].get("difficulty", "Medium") for i in pick]
                counts = {lvl: picked_lvls.count(lvl) for lvl in lvls}
                return counts
            counts_mcq = random_counts(mcq, num_mcq)
            counts_essay = random_counts(essay, num_essay)
        level_names = {"Easy": "DỄ", "Medium": "TRUNG BÌNH", "Hard": "KHÓ"}
        def check_and_warn(pool, counts, name):
            lvls = ["Easy", "Medium", "Hard"]
            pool_map = {lvl: [q for q in pool if q.get("difficulty") == lvl] for lvl in lvls}
            for lvl in lvls:
                required = counts[lvl]
                actual = len(pool_map[lvl])
                if actual < required:
                    msg = f"{name}: Không đủ số câu mức {level_names[lvl]} (cần {required}, có {actual}). Chỉ lấy được {actual} câu!"
                    warnings.append(msg)
                    counts[lvl] = actual
            return counts
        counts_mcq = check_and_warn(mcq, counts_mcq, "Trắc nghiệm")
        counts_essay = check_and_warn(essay, counts_essay, "Tự luận")
        counts = {"mcq": counts_mcq, "essay": counts_essay}

        # ------------------- ĐẢM BẢO RANDOM CÂU HỎI CHỈ 1 LẦN CHO MỖI MÃ ĐỀ ----------------
        exam_versions = {}  # key: code, value: {'mcq': [...], 'essay': [...]}

        for i in range(versions):
            code = string.ascii_uppercase[i]
            mcq_by_level = {level: [q for q in mcq if q.get("difficulty") == level] for level in ["Easy", "Medium", "Hard"]}
            essay_by_level = {level: [q for q in essay if q.get("difficulty") == level] for level in ["Easy", "Medium", "Hard"]}

            selected_mcq = []
            for level in ["Easy", "Medium", "Hard"]:
                n = counts["mcq"].get(level, 0)
                pool = mcq_by_level[level][:]
                if n > 0 and len(pool) >= n:
                    sampled = random.sample(pool, n)
                    selected_mcq.extend(sampled)
            random.shuffle(selected_mcq)

            selected_essay = []
            for level in ["Easy", "Medium", "Hard"]:
                n = counts["essay"].get(level, 0)
                pool = essay_by_level[level][:]
                if n > 0 and len(pool) >= n:
                    sampled = random.sample(pool, n)
                    selected_essay.extend(sampled)
            random.shuffle(selected_essay)

            exam_versions[code] = {"mcq": selected_mcq, "essay": selected_essay}

        # ------------------- SINH FILE .docx ĐỀ CŨNG NHƯ OFFICIAL TỪ BỘ CÂU HỎI ĐÃ RANDOM -----------------

        for i in range(versions):
            code = string.ascii_uppercase[i]
            selected_mcq = exam_versions[code]['mcq']
            selected_essay = exam_versions[code]['essay']

            # 1. Sinh exam_{code}.docx theo format cũ
            exam_file = generate_word_exams_from_selected(
                topic_path,
                selected_mcq,
                selected_essay,
                points_mcq,
                points_essay,
                code
            )
            files.append(exam_file)

            # 2. Sinh đề official & đáp án hội đồng giống hệt
            exam_fn = f"exam_{code}_official.docx"
            answer_fn = f"answer_{code}_official.docx"
            export_official_exam_format(
                topic_path,
                selected_mcq, selected_essay,
                num_mcq, num_essay, points_mcq, points_essay,
                file_name=exam_fn,
                subject=topic,
                code=code
            )
            export_official_answer(
                topic_path,
                selected_mcq, selected_essay,
                num_mcq, num_essay, points_mcq, points_essay,
                file_name=answer_fn,
                subject=topic,
                code=code
            )
            files.extend([exam_fn, answer_fn])
    return render_template(
        "form_exam.html",
        topics=topics,
        files=files,
        current_topic=current_topic,
        warnings=warnings,
    )

@app.route("/api/get_questions/<topic>")
def get_questions(topic):
    topic_path = os.path.join(UPLOAD_FOLDER, topic)
    bank = os.path.join(topic_path, "question_bank.json")
    if not os.path.exists(bank):
        return jsonify([])
    with open(bank, encoding="utf-8") as f:
        data = json.load(f)
    return jsonify(data.get("questions", []))

@app.route("/download/<topic>/<filename>")
def download(topic, filename):
    return send_from_directory(
        os.path.join(UPLOAD_FOLDER, topic, "exam"),
        filename,
        as_attachment=True
    )

@app.route("/api/delete_question/<topic>/<int:qidx>", methods=["POST"])
def delete_question(topic, qidx):
    topic_path = os.path.join(UPLOAD_FOLDER, topic)
    bank_file = os.path.join(topic_path, "question_bank.json")
    if not os.path.isfile(bank_file): return jsonify(success=False)
    with open(bank_file,encoding="utf-8") as f:
        data = json.load(f)
    qs = data.get("questions", [])
    if 0 <= qidx < len(qs):
        qs.pop(qidx)
        with open(bank_file,"w",encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return jsonify(success=True)
    return jsonify(success=False)

@app.route("/api/edit_question/<topic>/<int:qidx>", methods=["POST"])
def edit_question(topic, qidx):
    topic_path = os.path.join(UPLOAD_FOLDER, topic)
    bank_file = os.path.join(topic_path, "question_bank.json")
    if not os.path.isfile(bank_file): return jsonify(success=False)
    with open(bank_file,encoding="utf-8") as f:
        data = json.load(f)
    qs = data.get("questions", [])
    req = request.get_json()
    if 0 <= qidx < len(qs):
        qs[qidx].update({k:v for k,v in req.items() if k!='_idx'})
        with open(bank_file,"w",encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return jsonify(success=True)
    return jsonify(success=False)

@app.route("/")
def home():
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)